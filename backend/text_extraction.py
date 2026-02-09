import io
import pdfplumber
from docx import Document
from typing import Union


def extract_text_from_pdf(file_content: bytes) -> str:
    """
    Extract text from PDF file using pdfplumber.
    
    Args:
        file_content: PDF file content as bytes
    
    Returns:
        Extracted text as string
    """
    try:
        text_parts = []
        with pdfplumber.open(io.BytesIO(file_content)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        
        return "\n".join(text_parts)
    except Exception as e:
        raise ValueError(f"Error extracting text from PDF: {str(e)}")


def extract_text_from_docx(file_content: bytes) -> str:
    """
    Extract text from DOCX file using python-docx.
    
    Args:
        file_content: DOCX file content as bytes
    
    Returns:
        Extracted text as string
    """
    try:
        doc = Document(io.BytesIO(file_content))
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return "\n".join(text_parts)
    except Exception as e:
        raise ValueError(f"Error extracting text from DOCX: {str(e)}")


def extract_text_from_file(file_content: bytes, filename: str) -> str:
    """
    Extract text from a resume file (PDF or DOCX).
    
    Args:
        file_content: File content as bytes
        filename: Original filename (used to determine file type)
    
    Returns:
        Extracted text as string
    
    Raises:
        ValueError: If file type is unsupported or extraction fails
    """
    filename_lower = filename.lower()
    
    if filename_lower.endswith('.pdf'):
        return extract_text_from_pdf(file_content)
    elif filename_lower.endswith('.docx') or filename_lower.endswith('.doc'):
        return extract_text_from_docx(file_content)
    else:
        raise ValueError(f"Unsupported file type: {filename}. Only PDF and DOCX files are supported.")

